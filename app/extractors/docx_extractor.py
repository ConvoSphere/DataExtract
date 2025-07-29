"""
Extraktor für DOCX-Dateien.
"""

import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.extractors.base import BaseExtractor
from app.models.schemas import ExtractedText, FileMetadata, StructuredData


class DOCXExtractor(BaseExtractor):
    """Extraktor für DOCX-Dateien."""
    
    def __init__(self):
        super().__init__()
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx ist nicht installiert. Installieren Sie es mit: pip install python-docx")
        
        self.supported_extensions = [".docx", ".doc"]
        self.supported_mime_types = [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ]
        self.max_file_size = 50 * 1024 * 1024  # 50MB
    
    def can_extract(self, file_path: Path, mime_type: str) -> bool:
        """Prüft, ob der Extraktor die DOCX-Datei verarbeiten kann."""
        return (
            file_path.suffix.lower() in self.supported_extensions or
            mime_type in self.supported_mime_types
        )
    
    def extract_metadata(self, file_path: Path) -> FileMetadata:
        """Extrahiert Metadaten aus der DOCX-Datei."""
        stat = file_path.stat()
        
        metadata = FileMetadata(
            filename=file_path.name,
            file_size=stat.st_size,
            file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_extension=file_path.suffix.lower(),
            created_date=datetime.fromtimestamp(stat.st_ctime),
            modified_date=datetime.fromtimestamp(stat.st_mtime)
        )
        
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            if core_props.title:
                metadata.title = core_props.title
            if core_props.author:
                metadata.author = core_props.author
            if core_props.subject:
                metadata.subject = core_props.subject
            if core_props.keywords:
                metadata.keywords = [k.strip() for k in core_props.keywords.split(',')]
            
            # Seitenanzahl schätzen (ungefähr)
            paragraph_count = len(doc.paragraphs)
            metadata.page_count = max(1, paragraph_count // 20)  # Grobe Schätzung
            
        except Exception:
            pass
        
        return metadata
    
    def extract_text(self, file_path: Path) -> ExtractedText:
        """Extrahiert Text aus der DOCX-Datei."""
        content = ""
        
        try:
            doc = Document(file_path)
            
            # Text aus allen Absätzen extrahieren
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # Text aus Tabellen extrahieren
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content += cell.text + "\t"
                    content += "\n"
            
        except Exception as e:
            raise Exception(f"DOCX-Extraktion fehlgeschlagen: {str(e)}")
        
        # Text bereinigen
        content = self._clean_text(content)
        
        # Statistiken berechnen
        word_count = len(content.split()) if content else 0
        character_count = len(content)
        
        return ExtractedText(
            content=content,
            word_count=word_count,
            character_count=character_count
        )
    
    def extract_structured_data(self, file_path: Path) -> StructuredData:
        """Extrahiert strukturierte Daten aus der DOCX-Datei."""
        tables = []
        headings = []
        lists = []
        
        try:
            doc = Document(file_path)
            
            # Dokument durchgehen und Struktur analysieren
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    # Absatz
                    paragraph = Paragraph(element, doc)
                    text = paragraph.text.strip()
                    
                    if text:
                        # Überschriften erkennen
                        if self._is_heading(paragraph):
                            headings.append({
                                "level": self._get_heading_level(paragraph),
                                "text": text,
                                "position": len(headings)
                            })
                        
                        # Listen erkennen
                        if self._is_list_item(text):
                            lists.append(text)
                
                elif isinstance(element, CT_Tbl):
                    # Tabelle
                    table = Table(element, doc)
                    table_data = self._extract_table_data(table)
                    if table_data:
                        tables.append(table_data)
            
        except Exception:
            pass
        
        return StructuredData(
            tables=tables,
            headings=headings,
            lists=[lists] if lists else None
        )
    
    def _clean_text(self, text: str) -> str:
        """Bereinigt den extrahierten Text."""
        # Überflüssige Whitespaces entfernen
        text = re.sub(r'\s+', ' ', text)
        # Zeilenumbrüche normalisieren
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        # Tabulatoren durch Leerzeichen ersetzen
        text = text.replace('\t', ' ')
        return text.strip()
    
    def _is_heading(self, paragraph: Paragraph) -> bool:
        """Prüft, ob ein Absatz eine Überschrift ist."""
        # Prüfe auf Überschriften-Styles
        if paragraph.style.name.startswith('Heading'):
            return True
        
        # Prüfe auf Formatierung (fett, größere Schrift)
        for run in paragraph.runs:
            if run.bold or run.font.size and run.font.size.pt > 12:
                return True
        
        # Prüfe auf kurze Zeilen ohne Punkt am Ende
        text = paragraph.text.strip()
        if len(text) < 100 and text and text[-1] not in '.,!?':
            return True
        
        return False
    
    def _get_heading_level(self, paragraph: Paragraph) -> int:
        """Ermittelt die Hierarchie-Ebene einer Überschrift."""
        style_name = paragraph.style.name
        
        # Aus Style-Namen extrahieren
        if 'Heading' in style_name:
            try:
                level = int(style_name.replace('Heading ', ''))
                return min(level, 6)  # Maximal 6 Ebenen
            except ValueError:
                pass
        
        # Fallback: Basierend auf Formatierung
        for run in paragraph.runs:
            if run.font.size and run.font.size.pt > 16:
                return 1
            elif run.font.size and run.font.size.pt > 14:
                return 2
            elif run.bold:
                return 3
        
        return 1
    
    def _is_list_item(self, text: str) -> bool:
        """Prüft, ob ein Text ein Listenelement ist."""
        # Nummerierte Listen
        if re.match(r'^\d+\.\s', text):
            return True
        
        # Aufzählungslisten
        if re.match(r'^[•\-*]\s', text):
            return True
        
        # Buchstaben-Listen
        if re.match(r'^[a-zA-Z]\.\s', text):
            return True
        
        return False
    
    def _extract_table_data(self, table: Table) -> Dict[str, Any]:
        """Extrahiert Daten aus einer Tabelle."""
        try:
            rows = []
            headers = []
            
            for i, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                
                if i == 0:
                    # Erste Zeile als Header behandeln
                    headers = row_data
                else:
                    rows.append(row_data)
            
            return {
                "headers": headers,
                "rows": rows,
                "row_count": len(rows),
                "column_count": len(headers) if headers else 0
            }
            
        except Exception:
            return {}